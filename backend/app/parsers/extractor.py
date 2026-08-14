import io
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree
from defusedxml import ElementTree as DefusedElementTree

import pdfplumber
from bs4 import BeautifulSoup
from docx import Document


@dataclass
class Chapter:
    title: str
    text: str


@dataclass
class BookContent:
    title: str
    chapters: list[Chapter]
    word_count: int


@dataclass
class _Line:
    """A single visual line of text on a page, with its dominant font size."""
    text: str
    font_size: float
    top: float


# A line is treated as a page number if it's just a number, optionally wrapped
# in common decorations (e.g. "12", "- 12 -", "Page 12", "12 | Book Title").
_PAGE_NUMBER_RE = re.compile(
    r"^\s*(?:page\s+)?[-–—\|]?\s*(?:\d{1,4}|[ivxlcdm]{1,7})\s*[-–—\|]?\s*$",
    re.IGNORECASE,
)


def _page_lines(page) -> list[_Line]:
    """Group a page's words into visual lines with a representative font size."""
    words = page.extract_words(extra_attrs=["size"])
    grouped: dict[float, list] = {}
    for w in words:
        top = round(w["top"], 1)
        grouped.setdefault(top, []).append(w)

    lines: list[_Line] = []
    for top in sorted(grouped):
        line_words = sorted(grouped[top], key=lambda w: w["x0"])
        text = " ".join(w["text"] for w in line_words).strip()
        if not text:
            continue
        font_size = max(w.get("size", 12) for w in line_words)
        lines.append(_Line(text=text, font_size=font_size, top=top))
    return lines


def _find_running_lines(pages_lines: list[list[_Line]]) -> set[str]:
    """Detect repeated headers/footers.

    A short line near the top or bottom of the page that recurs on many pages is
    almost always a running header/footer (book title, chapter name, section),
    not real content — so it should not be narrated.
    """
    if len(pages_lines) < 4:
        return set()

    top_counts: Counter = Counter()
    bottom_counts: Counter = Counter()
    for lines in pages_lines:
        if not lines:
            continue
        # First two and last two lines are the header/footer candidates.
        for ln in lines[:2]:
            if len(ln.text) < 100:
                top_counts[ln.text] += 1
        for ln in lines[-2:]:
            if len(ln.text) < 100:
                bottom_counts[ln.text] += 1

    # Appears on at least ~40% of pages → it's boilerplate.
    threshold = max(3, int(len(pages_lines) * 0.4))
    running = set()
    for text, count in (top_counts + bottom_counts).items():
        if count >= threshold:
            running.add(text)
    return running


def _pdf_outline_titles(pdf) -> list[str]:
    """Return chapter titles from the PDF's table of contents / outline, if any."""
    try:
        outline = pdf.doc.get_outlines()
    except Exception:
        return []

    titles: list[str] = []
    for _level, title, *_ in outline:
        if title:
            cleaned = " ".join(str(title).split())
            if cleaned:
                titles.append(cleaned)
    return titles


def _ocr_space_pdf(file_bytes: bytes) -> str | None:
    """
    Best-effort OCR for scanned/image PDFs via OCR.Space, used only when normal
    text extraction comes up empty and OCR_SPACE_API_KEY is set. Note the free
    tier has small file-size / page limits, so this rescues small scanned docs;
    very large scanned books may exceed the free quota. Returns text or None.
    """
    import os

    key = os.environ.get("OCR_SPACE_API_KEY")
    if not key:
        return None
    try:
        import httpx

        resp = httpx.post(
            "https://api.ocr.space/parse/image",
            data={"apikey": key, "filetype": "PDF", "OCREngine": "2", "scale": "true"},
            files={"file": ("document.pdf", file_bytes, "application/pdf")},
            timeout=120.0,
        )
        resp.raise_for_status()
        data = resp.json()
        if data.get("IsErroredOnProcessing"):
            return None
        results = data.get("ParsedResults") or []
        text = "\n".join(r.get("ParsedText", "") for r in results).strip()
        return text or None
    except Exception:
        return None


def extract_from_pdf(file_bytes: bytes) -> BookContent:
    pdf = pdfplumber.open(io.BytesIO(file_bytes))
    title = (pdf.metadata or {}).get("Title", "") or "Untitled"

    pages_lines = [_page_lines(page) for page in pdf.pages]
    running = _find_running_lines(pages_lines)
    outline_titles = {t.lower() for t in _pdf_outline_titles(pdf)}

    chapters: list[Chapter] = []
    current_text = ""
    current_title = "Chapter 1"

    for lines in pages_lines:
        for ln in lines:
            text = ln.text

            # Drop running headers/footers and standalone page numbers.
            if text in running or _PAGE_NUMBER_RE.match(text):
                continue

            # Chapter break: either a large heading, or a line that matches a
            # title from the PDF's own table of contents.
            is_heading = ln.font_size > 16 and len(text) < 100
            is_toc_title = text.lower() in outline_titles and len(text) < 100

            if (is_heading or is_toc_title) and text:
                if current_text.strip():
                    chapters.append(Chapter(title=current_title, text=current_text.strip()))
                current_title = text
                current_text = ""
            else:
                current_text += text + " "

    if current_text.strip():
        chapters.append(Chapter(title=current_title, text=current_text.strip()))

    if not chapters:
        full_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        chapters = [Chapter(title="Full Text", text=full_text)]

    pdf.close()
    word_count = sum(len(ch.text.split()) for ch in chapters)

    # A scanned/image PDF has no text layer, so extraction yields ~nothing.
    # Fall back to OCR (if configured) to read the text off the page images.
    if word_count < 20:
        ocr_text = _ocr_space_pdf(file_bytes)
        if ocr_text and len(ocr_text.split()) > word_count:
            chapters = [Chapter(title=("Full Text" if title == "Untitled" else title), text=ocr_text)]
            word_count = len(ocr_text.split())

    return BookContent(title=title, chapters=chapters, word_count=word_count)


def _epub_spine_hrefs(zf: zipfile.ZipFile) -> tuple[str, list[str]]:
    """Parse the OPF to get the book title and spine-ordered content hrefs."""
    container = DefusedElementTree.fromstring(zf.read("META-INF/container.xml"))
    ns = {"c": "urn:oasis:names:tc:opendocument:xmlns:container"}
    rootfile_path = container.find(".//c:rootfile", ns).get("full-path")

    opf = DefusedElementTree.fromstring(zf.read(rootfile_path))
    opf_ns = {"opf": "http://www.idpf.org/2007/opf", "dc": "http://purl.org/dc/elements/1.1/"}
    opf_dir = rootfile_path.rsplit("/", 1)[0] + "/" if "/" in rootfile_path else ""

    title_el = opf.find(".//dc:title", opf_ns)
    title = title_el.text.strip() if title_el is not None and title_el.text else "Untitled"

    manifest = {}
    for item in opf.findall(".//opf:manifest/opf:item", opf_ns):
        manifest[item.get("id")] = opf_dir + item.get("href")

    spine_hrefs = []
    for itemref in opf.findall(".//opf:spine/opf:itemref", opf_ns):
        idref = itemref.get("idref")
        if idref in manifest:
            spine_hrefs.append(manifest[idref])

    return title, spine_hrefs


def extract_from_epub(file_bytes: bytes) -> BookContent:
    zf = zipfile.ZipFile(io.BytesIO(file_bytes))
    title, spine_hrefs = _epub_spine_hrefs(zf)

    chapters: list[Chapter] = []
    chapter_num = 1

    for href in spine_hrefs:
        try:
            content = zf.read(href)
        except KeyError:
            continue

        soup = BeautifulSoup(content, "html.parser")

        heading = soup.find(re.compile(r"^h[1-3]$"))
        chapter_title = heading.get_text(strip=True) if heading else f"Chapter {chapter_num}"

        paragraphs = soup.find_all("p")
        text = "\n".join(p.get_text(strip=True) for p in paragraphs)

        if text.strip():
            chapters.append(Chapter(title=chapter_title, text=text.strip()))
            chapter_num += 1

    zf.close()

    if not chapters:
        chapters = [Chapter(title="Full Text", text="No readable content found.")]

    word_count = sum(len(ch.text.split()) for ch in chapters)
    return BookContent(title=title, chapters=chapters, word_count=word_count)


def extract_from_docx(file_bytes: bytes) -> BookContent:
    doc = Document(io.BytesIO(file_bytes))
    title = doc.core_properties.title or "Untitled"

    chapters: list[Chapter] = []
    current_text = ""
    current_title = "Chapter 1"
    chapter_num = 1

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_text.strip():
                chapters.append(Chapter(title=current_title, text=current_text.strip()))
            chapter_num += 1
            current_title = para.text.strip() or f"Chapter {chapter_num}"
            current_text = ""
        else:
            current_text += para.text + "\n"

    if current_text.strip():
        chapters.append(Chapter(title=current_title, text=current_text.strip()))

    if not chapters:
        full_text = "\n".join(p.text for p in doc.paragraphs)
        chapters = [Chapter(title="Full Text", text=full_text)]

    word_count = sum(len(ch.text.split()) for ch in chapters)
    return BookContent(title=title, chapters=chapters, word_count=word_count)


def extract_from_txt(file_bytes: bytes) -> BookContent:
    text = file_bytes.decode("utf-8", errors="replace")

    chapter_pattern = re.compile(
        r"^(chapter\s+\d+[^\n]*|part\s+\d+[^\n]*)", re.IGNORECASE | re.MULTILINE
    )
    splits = chapter_pattern.split(text)

    chapters: list[Chapter] = []
    if len(splits) > 1:
        if splits[0].strip():
            chapters.append(Chapter(title="Introduction", text=splits[0].strip()))
        for i in range(1, len(splits), 2):
            title = splits[i].strip()
            body = splits[i + 1].strip() if i + 1 < len(splits) else ""
            if body:
                chapters.append(Chapter(title=title, text=body))
    else:
        chapters = [Chapter(title="Full Text", text=text.strip())]

    word_count = sum(len(ch.text.split()) for ch in chapters)
    return BookContent(title="Untitled", chapters=chapters, word_count=word_count)


EXTRACTORS = {
    ".pdf": extract_from_pdf,
    ".epub": extract_from_epub,
    ".docx": extract_from_docx,
    ".txt": extract_from_txt,
}


def extract_text(filename: str, file_bytes: bytes) -> BookContent:
    ext = Path(filename).suffix.lower()
    extractor = EXTRACTORS.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file format: {ext}")
    result = extractor(file_bytes)
    # If the extractor couldn't determine a title, use the filename stem
    if result.title == "Untitled":
        result.title = Path(filename).stem
    return result
